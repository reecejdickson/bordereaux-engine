"""Query pack routes - generate documents to send to MGAs"""
import uuid
from typing import Optional, List
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from pydantic import BaseModel
import io

from app.api.deps import get_tenant_context, TenantContext
from app.db.models import Issue, ProcessingRun

router = APIRouter()

# In-memory storage for query packs (would be DB in production)
QUERY_PACKS = {}

class QueryPackCreate(BaseModel):
    title: str
    issue_ids: List[str]
    include_evidence: bool = True

class QueryPackResponse(BaseModel):
    id: str
    title: str
    content: str
    issue_ids: List[str]
    created_at: str

@router.get("")
async def list_query_packs(run_id: Optional[str] = None, ctx: TenantContext = Depends(get_tenant_context)):
    packs = list(QUERY_PACKS.values())
    if run_id:
        packs = [p for p in packs if p.get("run_id") == run_id]
    return packs

@router.post("", response_model=QueryPackResponse)
async def create_query_pack(data: QueryPackCreate, ctx: TenantContext = Depends(get_tenant_context)):
    # Fetch issues
    issues = []
    for issue_id in data.issue_ids:
        result = await ctx.db.execute(select(Issue).where(Issue.id == uuid.UUID(issue_id)))
        issue = result.scalar()
        if issue:
            issues.append(issue)
    
    if not issues:
        raise HTTPException(status_code=400, detail="No valid issues found")
    
    # Generate content
    content = generate_query_content(data.title, issues, data.include_evidence)
    
    pack_id = str(uuid.uuid4())
    pack = {
        "id": pack_id,
        "title": data.title,
        "content": content,
        "issue_ids": data.issue_ids,
        "created_at": datetime.utcnow().isoformat(),
    }
    QUERY_PACKS[pack_id] = pack
    
    return QueryPackResponse(**pack)

@router.get("/{pack_id}")
async def get_query_pack(pack_id: str, ctx: TenantContext = Depends(get_tenant_context)):
    if pack_id not in QUERY_PACKS:
        raise HTTPException(status_code=404, detail="Query pack not found")
    return QUERY_PACKS[pack_id]

@router.get("/{pack_id}/export")
async def export_query_pack(pack_id: str, format: str = "markdown", ctx: TenantContext = Depends(get_tenant_context)):
    if pack_id not in QUERY_PACKS:
        raise HTTPException(status_code=404, detail="Query pack not found")
    
    pack = QUERY_PACKS[pack_id]
    content = pack["content"]
    
    if format == "docx":
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(pack["title"], 0)
            for para in content.split('\n\n'):
                doc.add_paragraph(para)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=query-pack.docx"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX export not available")
    
    elif format == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, height - 72, pack["title"])
            
            c.setFont("Helvetica", 11)
            y = height - 120
            for line in content.split('\n'):
                if y < 72:
                    c.showPage()
                    y = height - 72
                c.drawString(72, y, line[:90])
                y -= 14
            
            c.save()
            buffer.seek(0)
            
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=query-pack.pdf"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF export not available")
    
    else:  # markdown
        return StreamingResponse(
            io.BytesIO(content.encode()),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=query-pack.md"}
        )


def generate_query_content(title: str, issues: List, include_evidence: bool) -> str:
    """Generate query pack content"""
    lines = [
        f"# {title}",
        f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        f"Total Issues: {len(issues)}",
        "",
        "---",
        "",
    ]
    
    # Group by severity
    by_severity = {"P0": [], "P1": [], "P2": []}
    for issue in issues:
        sev = issue.severity.value if hasattr(issue.severity, 'value') else str(issue.severity)
        if sev in by_severity:
            by_severity[sev].append(issue)
    
    for sev in ["P0", "P1", "P2"]:
        if by_severity[sev]:
            sev_label = {"P0": "Critical", "P1": "Warning", "P2": "Info"}[sev]
            lines.append(f"## {sev} - {sev_label} ({len(by_severity[sev])} issues)")
            lines.append("")
            
            for i, issue in enumerate(by_severity[sev], 1):
                lines.append(f"### {i}. {issue.title}")
                lines.append(f"**Rule:** {issue.rule_id}")
                lines.append(f"**Description:** {issue.description}")
                
                if include_evidence and issue.cell_reference:
                    lines.append(f"**Cell Reference:** {issue.cell_reference}")
                    if issue.actual_value:
                        lines.append(f"**Found:** {issue.actual_value}")
                    if issue.expected_value:
                        lines.append(f"**Expected:** {issue.expected_value}")
                
                lines.append("")
            
            lines.append("---")
            lines.append("")
    
    lines.append("Please review and respond to each query above.")
    
    return '\n'.join(lines)
